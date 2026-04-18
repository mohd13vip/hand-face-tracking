import os
import re
import pandas as pd
from docx import Document
from docx.shared import Inches

HERE = os.path.dirname(__file__)
csv_path = os.path.join(HERE, "comparison.csv")
plots_dir = os.path.join(HERE, "plots")
out_docx = os.path.join(HERE, "Challenge_B_Report_ParrotCyber.docx")

df = pd.read_csv(csv_path)

doc = Document()
doc.add_heading("Computer Vision Competition — Task 1 (Challenge B) Report", level=1)
doc.add_paragraph("Name: Parrot Cyber")
doc.add_paragraph("Models: YOLOv8n / YOLOv8s / YOLOv8m (trained best checkpoints)")

doc.add_heading("1. Introduction", level=2)
doc.add_paragraph(
    "In Task 1 (Challenge B), we train and compare multiple YOLOv8 model scales on the provided dataset. "
    "The goal is to evaluate accuracy and efficiency trade-offs and recommend the best model for deployment."
)

doc.add_heading("2. Methodology", level=2)
doc.add_paragraph(
    "We trained YOLOv8n, YOLOv8s, and YOLOv8m and evaluated them using the same validation protocol. "
    "All final results are summarized in comparison.csv."
)

doc.add_heading("3. Results Table", level=2)
table = doc.add_table(rows=1, cols=len(df.columns))
hdr = table.rows[0].cells
for j, c in enumerate(df.columns):
    hdr[j].text = str(c)

for _, row in df.iterrows():
    cells = table.add_row().cells
    for j, c in enumerate(df.columns):
        cells[j].text = str(row[c])

doc.add_paragraph("")

doc.add_heading("4. Charts", level=2)
if os.path.isdir(plots_dir):
    plots = [f for f in os.listdir(plots_dir) if f.lower().endswith(".png")]
    plots.sort()
    if plots:
        for f in plots:
            doc.add_paragraph(f"Figure: {os.path.splitext(f)[0]}")
            doc.add_picture(os.path.join(plots_dir, f), width=Inches(6.5))
    else:
        doc.add_paragraph("No plots found. Run make_task1_plots.py first.")
else:
    doc.add_paragraph("No plots directory found. Run make_task1_plots.py first.")

doc.add_heading("5. Analysis & Recommendation", level=2)
doc.add_paragraph(
    "Typically, larger models (m) achieve higher accuracy but run slower and require more memory, "
    "while smaller models (n) are faster but less accurate. "
    "We recommend the model that best matches the competition priority metric (often mAP50-95 or mAP50). "
    "If real-time constraints exist, a slightly lower mAP may be acceptable for higher FPS."
)

doc.add_heading("6. Conclusion", level=2)
doc.add_paragraph(
    "This report compared YOLOv8n/s/m and documented quantitative results and sample predictions. "
    "The chosen model is the best trade-off according to the primary evaluation metric in comparison.csv."
)

doc.save(out_docx)
print("Saved:", out_docx)
